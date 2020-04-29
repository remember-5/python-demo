# -*- coding: utf-8 -*-
import os
from urllib.error import HTTPError

from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from docx.shared import Inches

# sohu主页
from com.remember.utils.db_utils import MyPymysqlPool
from com.remember.utils.html_utils import HtmlUtils
from com.remember.utils.image_utils import ImageUtils
from com.remember.utils.string_utils import StringUtils

SOHU = 'http://www.sohu.com'
# sohu自媒体账号页面 后面需要拼接author_id
SOHU_XPT = 'http://mp.sohu.com/profile?xpt='
# 图片保存路径
SAVE_PATH = 'D:\\file\\drawler\\sohu\\'
# mysql配置路径
CONF_PATH = os.path.join(os.path.dirname(os.getcwd()) + '\\config\\dbMysqlConfig.cnf')
# mysql实例
mysql = MyPymysqlPool(conf_name="dbMysql", conf_path=CONF_PATH)


class SohuUtils:

    @staticmethod
    def get_column1(html):
        """
        获取一级栏目列表
        :param html: bs4格式的html
        """
        column1 = []
        uls = html.find('nav', 'nav area').find_all('ul')
        for ul in iter(uls):
            col = {}
            lis = ul.find_all('li')
            for li in iter(lis):
                a_href = li.find('a').get('href')
                if 'http' not in a_href:
                    a_href = 'http{}'.format(a_href)
                col[li.text] = a_href
            column1.append(col)
        return column1

    @staticmethod
    def sohu_fashion():
        """
        获取时尚栏目下的二级栏目的所有用户
        :return:
        """
        # 一级栏目地址
        sohu_fashion_url = 'http://fashion.sohu.com'
        fhtml = HtmlUtils.get(sohu_fashion_url)
        lis = fhtml.find('ul', 'z-c-nav theme_fashion_border-top-color').find_all("li")
        fashion_columns = {}
        for li in iter(lis):
            a_href = li.find('a').get('href')
            if 'http' not in a_href:
                a_href = 'http:{}'.format(a_href)
                if '//' not in a_href:
                    a_href = sohu_fashion_url
            fashion_columns[li.text] = a_href
        # 删除首页
        del fashion_columns['首页']
        del fashion_columns['头排客']
        try:
            for s in fashion_columns:
                sz_url = fashion_columns[s]
                news = HtmlUtils.roll_html(sz_url, 2).find('div', 'news-wrapper').find_all(
                    attrs={'data-role': 'news-item'})
                _authors = {}
                for new in iter(news):
                    # 获取作者相关信息
                    author_span = new.find('div', class_='other').find('span', class_='name')
                    # 获取url
                    author_url = author_span.find('a').get('href')
                    # 获取url参数
                    url_params = HtmlUtils.get_url_param(author_url)
                    # author在sohu中的id
                    auther_id = url_params['xpt']
                    # 添加至全局
                    # author_dict[author_span.text] = auther_id
                    _authors[author_span.text] = auther_id
                SohuUtils.add_user(_authors, column1='时尚', column2=s)
            mysql.dispose()
        except Exception as e:
            print('异常',e)

    @staticmethod
    def user_article_list(user_name, user_url):
        """
        获取用户的所有文章列表
        :param user_name:用户名称
        :param user_url:用户地址
        :return:
        """
        # 获取用户的所有文章
        articles = HtmlUtils.roll_html(user_url, 2).find('ul', 'feed-list-area').find_all('li')
        print('[', user_name, ']', '文章条目=', len(articles))
        for article in iter(articles):
            # 获取用户单个文章的地址
            article_url = article.find('a').get("href")
            if 'http' not in article_url:
                article_url = 'http:{}'.format(article_url)
            # 获取单个文章的html
            article_html = HtmlUtils.get(article_url)
            # 交给处理类来处理
            SohuUtils.sohu_article(article_html)

    @staticmethod
    def sohu_article(child_html):
        # 文章标题
        title = child_html.find(attrs={'data-role': 'original-title'}).text
        print('文章标题=', title)
        # 所有信息
        article = child_html.find('article', class_='article')
        if article is None:
            print('文章标题=', title, '文章不存在或异常')
            return
        ps = article.find_all('p')

        document = Document()
        for p in iter(ps):
            # 查找是否包含img标签
            is_img = p.find('img')
            # 文章
            if is_img is None:
                b = p.text
                if b is None:
                    b = p.find('span').text
                document.add_paragraph(b)
            else:
                image_path = p.find('img').get("src")
                ht = StringUtils.find_last_index(image_path, "http:")
                if -1 == ht:
                    image_path = 'http:{}'.format(image_path)
                image_name_start = StringUtils.find_last_index(image_path, "/")
                image_name = image_path[image_name_start + 1:]
                try:
                    print('imageName=', image_name, 'imagePath=http:', image_path)
                    ImageUtils.save_web_image(image_path, SAVE_PATH + image_name)
                except HTTPError:
                    print('error')
                    break
                image_width = ImageUtils.web_image_width(image_path)
                # word 保存图片
                _width = Inches(image_width / 300)
                try:
                    document.add_picture(SAVE_PATH + image_name, width=_width)
                except (UnrecognizedImageError, ZeroDivisionError, Exception):
                    print('图片获取异常')
                    pass
        title = title.replace('/', "").replace('\\', "").replace(':', ""). \
            replace('*', "").replace('?', "").replace('<', "").replace('>', "").replace('|', "")
        document.save(SAVE_PATH + title[10:] + '.docx')  # 保存文档

    @staticmethod
    def add_column1():
        sohu_html = HtmlUtils.get(SOHU)
        column1 = SohuUtils.get_column1(sohu_html)
        """
        添加一级栏目
        :return:
        """
        for col in column1:
            _index = 0
            for c in col:
                sql = """INSERT INTO sh_column (name,url,level,parentid,spm) VALUES(%s,%s,%s,%s,%s)"""
                param = (c, col[c], 1, 0, _index)
                _index += 1
                mysql.insert(sql, param)
        mysql.dispose()

    @staticmethod
    def add_user(_authors=None, column1=None, column2=None):
        """
        添加用户
        :return:
        """
        print('1')
        for user in iter(_authors):
            select_count = "SELECT count(1) count FROM sh_user WHERE name = '%s'" % user
            select_result = mysql.get_one(select_count)
            isexist = select_result['count']
            if 0 == isexist:
                # 可以插入
                query_col1 = """SELECT id FROM sh_column WHERE name = '%s'""" % column1
                column1_id = mysql.get_one(query_col1)['id']
                query_col2 = """SELECT id FROM sh_column WHERE name = '%s'""" % column2
                column2_id = mysql.get_one(query_col2)['id']
                code = _authors[user]
                insert_sql = """INSERT INTO sh_user(name,code,address,column1,column2) VALUES('%s','%s','%s',%d,%d)""" % (
                    user, code, SOHU_XPT + code, column1_id, column2_id)
                print(insert_sql)
                mysql.insert(insert_sql)


if __name__ == '__main__':
    # sohu_html = HtmlUtils.get(SOHU)
    # SohuUtils.get_column1(sohu_html)
    SohuUtils.sohu_fashion()
# print(author_dict)

# for (k, v) in author_dict.items():
#     SohuUtils.user_article_list(k, SOHU_XPT + v)

# SohuUtils.sohu_article(HtmlUtils.get('http://www.sohu.com/a/313303133_377096'))
# SohuUtils.sohu_article(HtmlUtils.get('http://www.sohu.com/a/312129934_220816'))
