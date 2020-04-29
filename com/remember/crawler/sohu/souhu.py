# -*- coding: utf-8 -*-
from urllib.error import HTTPError

from docx import Document
from docx.shared import Inches

# shUrl = 'https://mp.sohu.com/profile?xpt=MzYwOTU4Mzg5NUBzaW5hLnNvaHUuY29t'
# shUrl = 'http://mp.sohu.com/profile?xpt=c29odW1wNnNsc2w0QHNvaHUuY29t'
from com.remember.utils.html_utils import HtmlUtils
from com.remember.utils.image_utils import ImageUtils
from com.remember.utils.string_utils import StringUtils

shUrl = 'http://mp.sohu.com/profile?xpt=MTE5ODk3NjY2MEBzaW5hLnNvaHUuY29t'
savePath = 'D:\\file\\drawler\\sohu\\'


class DisposeChild:
    # 处理文章
    @staticmethod
    def dis(child_html):
        # 文章标题
        title = child_html.find(attrs={'data-role': 'original-title'}).text
        print(title)
        # 所有信息
        ps = child_html.find('article', class_='article').find_all('p')
        document = Document()
        for p in iter(ps):
            # 查找是否包含img标签
            is_img = p.find('img')
            # 文章
            if is_img is None:
                b = p.text
                if b is None:
                    b = p.find('span').text
                document.add_paragraph(b)
            else:
                image_path = p.find('img').get("src")
                ht = StringUtils.find_last_index(image_path, "http:")
                if -1 == ht:
                    image_path = 'http:{}'.format(image_path)
                image_name_start = StringUtils.find_last_index(image_path, "/")
                image_name = image_path[image_name_start + 1:]
                try:
                    print('imageName=', image_name, 'imagePath=http:', image_path)
                    ImageUtils.save_web_image(image_path, savePath + image_name)
                except HTTPError:
                    print('error')
                    break
                image_width = ImageUtils.web_image_width(image_path)
                # word 保存图片
                document.add_picture(savePath + image_name, width=Inches(int(image_width) / 300))
        document.save(savePath + title[10:] + '.docx')  # 保存文档


if __name__ == '__main__':
    articles = HtmlUtils.get(shUrl).find('ul', 'feed-list-area').find_all('li')
    print('文章条目=', len(articles))
    article = iter(articles)
    for x in article:
        url = 'http:' + x.find('a').get("href")
        print(url)
        childHtml = HtmlUtils.get(url)
        DisposeChild.dis(childHtml)
