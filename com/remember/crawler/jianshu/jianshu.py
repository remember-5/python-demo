# -*- coding: utf-8 -*-
from urllib.error import HTTPError
from docx import Document
from docx.shared import Inches

from com.remember.utils.html_utils import HtmlUtils
from com.remember.utils.image_utils import ImageUtils
from com.remember.utils.string_utils import StringUtils

cUrl = ['https://www.jianshu.com/u/7ec548858a10', 'https://www.jianshu.com/c/4182da9b406b']
prefixUrl = 'https://www.jianshu.com/'
wordPath = 'D:\\file\\drawler\\zlq\\'


class DisposeChild:
    @staticmethod
    def dis(child_html):
        # 判断是否传空
        if child_html is None:
            return None
        # 文章详情
        article = child_html.find('div', class_='article')
        # 标题（word文件名）
        title = article.find('h1', class_='title').text
        # 正文
        content = article.find('div', class_='show-content-free')
        # 正文所有p div标签（因为有图片 所以要解析）
        all_tag = content.find_all(['p', 'div'])
        # 迭代器遍历
        tags = iter(all_tag)

        print("文章=", title)

        # 保存到word
        document = Document()

        for tag in tags:
            # 判断tag为p的话是文章，直接添加到word
            tag_name = tag.name
            if 'p' == tag_name:
                document.add_paragraph(tag.text)
            elif 'div' == tag_name:
                # 判断可能是图片
                clazz = tag.get('class')[0]
                # 根据class来判断是不是图片
                if clazz == 'image-view':
                    # 文章插入图片
                    # 图片地址
                    image_path = tag.next.get('data-original-src')
                    # 图片宽度（需要转英寸，html拿到的是像素  300像素/英寸）
                    image_width = tag.next.get('data-original-width')
                    # 保存图片到本地
                    # 判断是否携带格式（带.jpeg等） 清空.后面的数据
                    net_count = StringUtils.find_last_index(image_path, '.jpeg')
                    if -1 != net_count:
                        image_path = image_path[0:net_count]
                    image_name = image_path[StringUtils.find_last_index(image_path, '/') + 1:]
                    print('imageName=', image_name, 'imagePath=http:', image_path)
                    try:
                        ImageUtils.save_web_image('http:' + image_path, wordPath + image_name + ".jpg")
                    except HTTPError:
                        print('error')
                        break
                    document.add_picture(wordPath + image_name + ".jpg", width=Inches(int(image_width) / 300))

        # 以标题命名word的名字会有异常字符，做替换
        title = title.replace('/', "").replace('\\', "").replace(':', ""). \
            replace('*', "").replace('?', "").replace('<', "").replace('>', "").replace('|', "")
        document.save(wordPath + title + '.docx')  # 保存文档


if __name__ == '__main__':
    for u in range(0, len(cUrl)):
        html = HtmlUtils.time_roll_html(url=cUrl[u])
        de = html.find_all("ul", class_='note-list')
        for d in range(0, len(de)):
            lis = de[d].find_all("li")
            li = iter(lis)
            print('长度=', len(lis))
            for x in li:
                us = x.find_all('a', class_='wrap-img')
                if len(us) == 1:
                    childUrl = us[0].get('href')
                    cHtml = HtmlUtils.get(prefixUrl + childUrl)
                    DisposeChild.dis(cHtml)
