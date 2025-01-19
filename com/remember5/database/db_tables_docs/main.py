import os
import re
import click
import pandas as pd
import sqlparse
import pymysql
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_sql_statements(sql_content: str) -> List[Dict[str, str]]:
    """解析 SQL 内容并返回表结构数据"""
    tables_data = []
    parsed = sqlparse.parse(sql_content)

    for statement in parsed:
        if statement.get_type() == 'CREATE':
            current_table = re.search(r'CREATE TABLE `?(\w+)`?', str(statement)).group(1)

            # 解析列定义
            columns = re.findall(
                r'`(\w+)`\s+([\w()]+)(?:\s+CHARACTER SET.*?)?(?:\s+COLLATE.*?)?(?:\s+NOT NULL)?(?:\s+DEFAULT\s+(.*?))?(?:\s+COMMENT\s+\'(.*?)\')?(?=,|\n|$)',
                str(statement)
            )

            for col in columns:
                column_name = col[0]
                if column_name == current_table:
                    continue  # 跳过表名行

                data_type = col[1]
                length = ''
                decimal = ''
                if '(' in data_type:
                    type_parts = data_type.split('(')
                    data_type = type_parts[0]
                    length_part = type_parts[1].replace(')', '')
                    if ',' in length_part:
                        length, decimal = length_part.split(',')
                    else:
                        length = length_part

                is_primary = 'YES' if f'PRIMARY KEY (`{column_name}`)' in str(statement) else 'NO'
                allow_null = 'NO' if f'`{column_name}` {data_type}' in str(statement) and 'NOT NULL' in str(statement) else 'YES'
                default_value = col[2] if col[2] else ''
                comment = col[3] if len(col) > 3 else ''

                tables_data.append({
                    '表名': current_table,
                    '字段名': column_name,
                    '数据类型': data_type,
                    '长度': length,
                    '小数位': decimal,
                    '允许空值': allow_null,
                    '主键': is_primary,
                    '默认值': default_value,
                    '说明': comment
                })

    return tables_data


def sql_to_excel(sql_file_path: str, output_excel_path: str) -> str:
    """解析 SQL 文件并生成 Excel 文档"""
    try:
        with open(sql_file_path, 'r', encoding='utf-8') as f:
            sql_content = f.read()

        tables_data = parse_sql_statements(sql_content)

        # 生成 Excel 文件
        df = pd.DataFrame(tables_data)
        df.to_excel(output_excel_path, index=False, engine='openpyxl')
        return f"✨ Excel 文档已生成: {output_excel_path}"
    except Exception as e:
        return f"❌ 生成 Excel 文档时出错: {e}"


def sql_to_docs(sql_file_path: str, output_doc_path: str) -> str:
    """解析 SQL 文件并生成 Word 文档"""
    try:
        with open(sql_file_path, 'r', encoding='utf-8') as f:
            sql_content = f.read()

        tables_data = parse_sql_statements(sql_content)

        # 创建 Word 文档
        document = Document()

        for table_name, group in pd.DataFrame(tables_data).groupby('表名'):
            # 添加表名
            table_heading = document.add_paragraph(f"表 {table_name}")
            table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 创建表格
            table = document.add_table(rows=1, cols=9)
            table.style = 'Table Grid'

            # 添加表头
            headers = ['表名', '字段名', '数据类型', '长度', '小数位', '允许空值', '主键', '默认值', '说明']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加表格数据
            for _, row in group.iterrows():
                cells = table.add_row().cells
                row_data = [
                    row['表名'], row['字段名'], row['数据类型'], row['长度'],
                    row['小数位'], row['允许空值'], row['主键'], row['默认值'], row['说明']
                ]
                for idx, value in enumerate(row_data):
                    cells[idx].text = str(value)
                    cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 设置列宽
            for cell in table.columns[0].cells:
                cell.width = Inches(0.5)

            # 添加间隔
            document.add_paragraph()

        # 保存文档
        document.save(output_doc_path)
        return f"✨ Word 文档已生成: {output_doc_path}"
    except Exception as e:
        return f"❌ 生成 Word 文档时出错: {e}"


def connect_to_database(host: str, user: str, password: str, database: str, port: str) -> pymysql.connections.Connection:
    """连接到 MySQL 数据库"""
    try:
        # 将端口号从字符串转换为整数
        try:
            port = int(port)
            if not (1 <= port <= 65535):
                raise ValueError("端口号必须在 1 到 65535 之间")
        except ValueError as e:
            raise ValueError(f"无效的端口号: {port}。{e}")

        db = pymysql.connect(host=host, user=user, password=password, database=database, port=port)
        print("✨ 数据库连接成功，准备开始提取数据...")
        return db
    except pymysql.Error as e:
        print(f"❌ 数据库连接失败: {e}")
        raise


def fetch_table_info(cursor: pymysql.cursors.Cursor, database: str) -> List[Dict[str, str]]:
    """获取所有表名和表注释"""
    cursor.execute("""
        SELECT TABLE_NAME, TABLE_COMMENT
        FROM information_schema.TABLES
        WHERE TABLE_SCHEMA = %s
    """, (database,))
    return [{"Table": table_name, "Table Comment": table_comment} for table_name, table_comment in cursor.fetchall()]


def fetch_column_info(cursor: pymysql.cursors.Cursor, table_name: str, database: str) -> List[Dict[str, Optional[str]]]:
    """获取表的字段信息"""
    cursor.execute("""
        SELECT COLUMN_NAME, COLUMN_TYPE, IS_NULLABLE, COLUMN_DEFAULT, COLUMN_COMMENT
        FROM information_schema.COLUMNS
        WHERE TABLE_NAME = %s AND TABLE_SCHEMA = %s
    """, (table_name, database))
    return [{
        "Field": column_name,
        "Type": column_type,
        "Null": is_nullable,
        "Default": column_default if column_default is not None else "",
        "Comment": column_comment if column_comment is not None else ""
    } for column_name, column_type, is_nullable, column_default, column_comment in cursor.fetchall()]


def generate_excel_from_db(host: str, user: str, password: str, database: str, port: str, output_path: str) -> str:
    """从数据库生成 Excel 文件"""
    try:
        db = connect_to_database(host, user, password, database, port)
        cursor = db.cursor()

        # 获取所有表信息
        table_info = fetch_table_info(cursor, database)

        # 获取所有字段信息
        all_columns = []
        for table in table_info:
            columns = fetch_column_info(cursor, table["Table"], database)
            for column in columns:
                all_columns.append({**table, **column})

        # 生成 Excel 文件
        df = pd.DataFrame(all_columns)
        df.to_excel(output_path, index=False)
        return f"✨ Excel 文档已生成: {output_path}"

    except Exception as e:
        return f"❌ 发生错误: {e}"
    finally:
        if 'db' in locals() and db.open:
            db.close()
            print("🔒 数据库连接已关闭。")


@click.group()
def cli():
    # 添加优雅的标题
    click.echo(click.style('''
    ╭──────────────────────────────╮
    │   SQL Table Document Maker   │
    ╰──────────────────────────────╯
    ''', fg='cyan'))
    """SQL 和数据库结构导出工具"""
    pass


@cli.command()
@click.option('--host', prompt='数据库主机', help='数据库主机地址')
@click.option('--user', prompt='数据库用户', help='数据库用户名')
@click.option('--password', prompt='数据库密码', hide_input=True, help='数据库密码')
@click.option('--database', prompt='数据库名称', help='数据库名称')
@click.option('--port', prompt='数据库端口', default='3306', help='数据库端口号')
@click.option('--output', default=None, help='输出 Excel 文件路径')
def db(host: str, user: str, password: str, database: str, port: str, output: str):
    """从数据库生成 Excel 文件"""
    if output is None:
        output = f"{database}.xlsx"
    result = generate_excel_from_db(host, user, password, database, port, output)
    click.echo(result)


@cli.command()
@click.option('--sql', prompt='SQL 文件路径', help='SQL 文件路径')
@click.option('--excel', default=None, help='输出 Excel 文件路径')
@click.option('--doc', default=None, help='输出 Word 文件路径')
def sql(sql: str, excel: Optional[str], doc: Optional[str]):
    """从 SQL 文件生成 Excel 和 Word 文档"""
    try:
        # 自动生成输出文件路径
        base_name = os.path.splitext(os.path.basename(sql))[0]
        if excel is None:
            excel = f"{base_name}.xlsx"
        if doc is None:
            doc = f"{base_name}.docx"

        # 生成 Excel 文档
        excel_result = sql_to_excel(sql, excel)
        click.echo(excel_result)

        # 生成 Word 文档
        doc_result = sql_to_docs(sql, doc)
        click.echo(doc_result)
    except Exception as e:
        click.echo(f"❌ 生成文档时出错: {e}")


if __name__ == "__main__":
    cli()
